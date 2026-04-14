"""
ergane/filters/cv_generator.py
Generate tailored CV and cover letter using Claude API.
Uses Anthropic's Claude Sonnet 4.5 for high-quality generation.
"""
import logging
import os
import re
from pathlib import Path
from typing import Optional, Tuple

import requests
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv

load_dotenv()

from db.models import Job

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")
ANTHROPIC_ENABLED = bool(ANTHROPIC_API_KEY)
ANTHROPIC_URL = "https://api.anthropic.com/v1/messages"
ANTHROPIC_MODEL = "claude-sonnet-4-20250514"

# Contact information
MAYTE_EMAIL = os.getenv("MAYTE_EMAIL", "")
MAYTE_PHONE = os.getenv("MAYTE_PHONE", "")
MAYTE_LINKEDIN = os.getenv("MAYTE_LINKEDIN", "")
MAYTE_GITHUB = os.getenv("MAYTE_GITHUB", "")

# ---------------------------------------------------------------------------
# Mayte's Base CV Template
# ---------------------------------------------------------------------------

def _build_cv_base() -> str:
    """Build CV base with contact info from environment."""
    contact_parts = ["Location: Ciudad de Mexico, Mexico"]
    if MAYTE_EMAIL:
        contact_parts.insert(0, f"Email: {MAYTE_EMAIL}")
    if MAYTE_PHONE:
        contact_parts.insert(1 if MAYTE_EMAIL else 0, f"Phone: {MAYTE_PHONE}")
    if MAYTE_LINKEDIN:
        contact_parts.append(f"LinkedIn: {MAYTE_LINKEDIN}")
    if MAYTE_GITHUB:
        contact_parts.append(f"GitHub: {MAYTE_GITHUB}")
    contact_line = " | ".join(contact_parts)

    return f"""
CANDIDATE: Mayte Giovanna Hernandez Rios
CONTACT: {contact_line}
PROFILE:
- Cloud & Automation Engineer, 1 year AWS enterprise support (S3, DataSync, Transfer Family, Lambda)
- Python/boto3 automation, Terraform IaC
- AI/ML: RAG, LangChain, Anthropic API, local inference (Ollama/ROCm)
- Math @ UNAM, building MLOps/LLMOps career
- Co-founder Positronica Labs (FairHire, QMANUS - production AI systems)

TECHNICAL SKILLS:
- Languages: Python, JavaScript, TypeScript, SQL
- Cloud/AWS: S3, Lambda, IAM, EC2, DataSync, Transfer Family, CloudWatch
- IaC/DevOps: Terraform, Docker, GitHub Actions, CI/CD
- AI/ML: LangChain, RAG, LLM ops, Ollama, Anthropic Claude API, Hugging Face
- Backend: FastAPI, Flask, SQLAlchemy, Pydantic, asyncio
- Frontend: React, TypeScript, Next.js, Tailwind CSS, shadcn/ui
- Databases: SQLite, PostgreSQL, MySQL, TiDB, vector databases (Pinecone, Chroma)
- Data: pandas, NumPy, Airflow, data pipelines, ETL
- Tools: Git, Linux, bash, Playwright, Scrapling

PROJECTS:
- Ergane: Automated job scraper with CV matching (Python, Scrapling, Telegram bot)
- FairHire: AI-powered job fit analyzer for women in Mexico (React, Claude API)
- QMANUS: Multi-tenant AI agent platform (FastAPI, TiDB Cloud, Qwen/DashScope)
- UIGen: AI-powered React component generator (Next.js, Vercel AI SDK)

EDUCATION:
- BSc Mathematics, UNAM (expected 2027)
- AWS Certified (in progress)

LANGUAGES:
- Spanish (native)
- English (C1 professional)
"""

MAYTE_CV_BASE = _build_cv_base()



# ---------------------------------------------------------------------------
# Language Detection
# ---------------------------------------------------------------------------

def _detect_language(text: str) -> str:
    """Detect if job posting is in Spanish or English."""
    spanish_markers = [
        'busco', 'buscamos', 'empresa', 'experiencia', 'requisitos',
        'habilidades', 'trabajo', 'remoto', 'vacante', 'postular',
        'equipo', 'ofrecemos', 'deseable', 'preferentemente', 'ingeniero',
        'desarrollador', 'estamos', 'queremos', 'bienvenido', 'salario',
        'sueldo', 'ubicacion', 'ubicación', 'eres', 'tienes', 'oferta',
        'contratamos', 'unete', 'únete'
    ]
    count = sum(1 for w in spanish_markers if w in text.lower())
    return "Spanish" if count >= 3 else "English"

# ---------------------------------------------------------------------------
# CV Generation Prompts
# ---------------------------------------------------------------------------

CV_GENERATION_PROMPT = """You are a professional CV writer specializing in tech roles for the Mexican/LATAM market.

TASK: Generate a tailored CV for Mayte Giovanna Hernandez Rios based on the job description provided.

CANDIDATE BASE CV:
{base_cv}

JOB DESCRIPTION:
Title: {job_title}
Company: {company}
Description: {description}

LANGUAGE: Write the entire CV in {language}. Match the language of the job posting exactly.

FORMAT STRUCTURE - Follow this layout exactly:

# MAYTE GIOVANNA HERNANDEZ RIOS

*Cloud & Automation Engineer*

## PROFESSIONAL SUMMARY
[3-4 sentences tailored to this job. No bullets.]

## TECHNICAL SKILLS
**Cloud & DevOps:** AWS (S3, Lambda, IAM, EC2, DataSync), Terraform, Docker, GitHub Actions, CI/CD
**AI/ML:** LangChain, RAG, Claude API, Ollama, Hugging Face, Embeddings, Prompt Engineering
**Backend:** Python, FastAPI, Flask, SQLAlchemy, Pydantic, asyncio
**Frontend:** React, TypeScript, Next.js, Tailwind CSS, shadcn/ui
**Databases:** PostgreSQL, SQLite, TiDB, Vector Databases (Pinecone, Chroma)
**Data & Tools:** pandas, NumPy, Airflow, ETL, Git, Linux, bash, Playwright

## PROFESSIONAL EXPERIENCE

### [Company/Project] | [Role] | Remote/Location | [Year(s)]
- Achievement with specific impact or metric
- Achievement with specific impact or metric

## PROJECTS

### [Project Name]
**Tech:** [Technologies used]
- What it does and impact

## EDUCATION

### BSc Mathematics, UNAM | Expected 2027
- Relevant coursework or achievements

## LANGUAGES
**Spanish:** Native
**English:** C1 Professional

FORMATTING RULES:
1. Use ## for all section headers (will be formatted as bold + underline in Word)
2. Use ### for subsection headers like company/project names
3. Use **bold:** for category labels in the skills section
4. Use - for all bullet points (no other list formats)
5. No tables, no columns, no multi-column layouts
6. No emojis, no special symbols, no unicode formatting
7. Contact info is injected separately — do NOT include email, phone, LinkedIn, or GitHub in the CV
8. Keep it to 1 page maximum (about 30-35 lines total)
9. All section headers must be in UPPERCASE
10. Do not add a date or timestamp

INSTRUCTIONS:
1. Highlight the most relevant skills from the base CV that match this job
2. Reorder sections to emphasize strengths for this specific role
3. Add a professional summary tailored to this job
4. Keep technical accuracy - don't claim skills Mayte doesn't have
5. Use action verbs and quantify achievements where possible
6. Languages section MUST show: Spanish (Native), English (C1 Professional). Do not change these levels.
7. Do not use emojis or special characters anywhere in the CV

OUTPUT FORMAT:
Return ONLY the CV in Markdown format. No additional text or explanations. Start with # and end with the Languages section."""

COVER_LETTER_PROMPT = """You are a professional cover letter writer for tech roles in Mexico/LATAM.

TASK: Write a compelling cover letter for Mayte Giovanna Hernandez Rios applying to this job.

CANDIDATE PROFILE:
- Cloud & Automation Engineer with 1 year AWS enterprise experience
- Python/boto3 automation expert, Terraform IaC
- AI/ML: RAG, LangChain, Anthropic API, local LLM inference
- Math @ UNAM, co-founder Positronica Labs
- Bilingual: Spanish (native), English (C1)

JOB DESCRIPTION:
Title: {job_title}
Company: {company}
Description: {description}

LANGUAGE: Write the entire cover letter in {language}. Match the language of the job posting exactly.

FORMAT:
- 3-4 paragraphs maximum
- First paragraph: Strong opening showing genuine interest in THIS company/role
- Middle paragraphs: 2-3 specific achievements that match the job requirements
- Final paragraph: Clear call to action (e.g., "I'm excited to discuss how I can contribute to your team")
- Professional but warm tone suited to Mexican business culture
- No emojis, no special characters, plain text and standard punctuation only

INSTRUCTIONS:
1. Demonstrate knowledge of the company if available from the job description
2. Connect your specific experience to their stated needs
3. Use action verbs and concrete examples
4. Show cultural fit and genuine enthusiasm
5. Keep it to 250-350 words

OUTPUT FORMAT:
Return ONLY the cover letter body (no salutation, no signature). No additional text or explanations."""

# ---------------------------------------------------------------------------
# Main Functions
# ---------------------------------------------------------------------------

def generate_cv(job: Job) -> Tuple[Optional[str], Optional[str]]:
    """
    Generate tailored CV and cover letter for a job.

    Args:
        job: Job object with description

    Returns:
        (cv_text, cover_letter_text) or (None, None) if generation fails
    """
    if not ANTHROPIC_ENABLED:
        logger.warning("Anthropic API not configured, skipping CV generation")
        return (None, None)

    try:
        # Prepare job context
        job_title = job.title or "Unknown position"
        company = job.company or "Unknown company"
        description = job.description or "No description available"

        # Truncate description if too long (max 3000 chars for context)
        if len(description) > 3000:
            description = description[:3000] + "..."

        # Detect job posting language
        language = _detect_language(description)
        logger.info("[%s] Detected job language: %s", job.source, language)

        # Generate CV
        cv_text = _call_claude(
            CV_GENERATION_PROMPT.format(
                base_cv=MAYTE_CV_BASE,
                job_title=job_title,
                company=company,
                description=description,
                language=language,
            )
        )

        if not cv_text:
            logger.error("[%s] CV generation failed", job.source)
            return (None, None)

        logger.info("[%s] CV generated successfully (%d chars)", job.source, len(cv_text))

        # Generate cover letter
        cover_letter = _call_claude(
            COVER_LETTER_PROMPT.format(
                job_title=job_title,
                company=company,
                description=description,
                language=language,
            )
        )

        if cover_letter:
            logger.info("[%s] Cover letter generated (%d chars)", job.source, len(cover_letter))
        else:
            logger.warning("[%s] Cover letter generation failed", job.source)

        return (cv_text, cover_letter)

    except Exception as e:
        logger.exception("[%s] Error generating CV: %s", job.source, e)
        return (None, None)


def _call_claude(prompt: str) -> Optional[str]:
    """
    Call Anthropic Claude API.

    Args:
        prompt: The prompt to send

    Returns:
        Generated text or None if failed
    """
    headers = {
        "x-api-key": ANTHROPIC_API_KEY,
        "Content-Type": "application/json",
        "anthropic-version": "2023-06-01",
    }

    payload = {
        "model": ANTHROPIC_MODEL,
        "max_tokens": 2048,
        "messages": [
            {"role": "user", "content": prompt}
        ]
    }

    try:
        response = requests.post(ANTHROPIC_URL, headers=headers, json=payload, timeout=60)
        response.raise_for_status()

        result = response.json()
        content = result.get("content", [])
        
        if content and len(content) > 0:
            return content[0].get("text", "")
        
        return None

    except requests.exceptions.RequestException as e:
        logger.error("Anthropic API call failed: %s", e)
        return None


def generate_cv_simple(job_description: str) -> Optional[str]:
    """
    Generate CV from a simple job description string (for Telegram bot).

    Args:
        job_description: Job description text

    Returns:
        Generated CV text or None
    """
    if not ANTHROPIC_ENABLED:
        logger.warning("Anthropic API not configured")
        return None

    language = _detect_language(job_description)

    prompt = CV_GENERATION_PROMPT.format(
        base_cv=MAYTE_CV_BASE,
        job_title="Position",
        company="Company",
        description=job_description,
        language=language,
    )

    return _call_claude(prompt)


def generate_cv_word(job: Job, output_path: Optional[str] = None) -> Optional[str]:
    """
    Generate tailored CV as a Word document (.docx).

    Args:
        job: Job object with description
        output_path: Optional path for output file. If None, saves to ./cv_output/{company}_{job_title}_{date}.docx

    Returns:
        Path to generated .docx file or None if generation fails
    """
    if not ANTHROPIC_ENABLED:
        logger.warning("Anthropic API not configured, skipping CV generation")
        return None

    try:
        # Generate CV markdown first
        cv_text, _ = generate_cv(job)
        if not cv_text:
            logger.error("[%s] CV generation failed", job.source)
            return None

        # Convert to Word
        docx_path = markdown_to_word(cv_text, output_path, job.company, job.title)
        logger.info("[%s] CV Word document generated: %s", job.source, docx_path)
        return docx_path

    except Exception as e:
        logger.exception("[%s] Error generating CV Word document: %s", job.source, e)
        return None


def markdown_to_word(markdown_text: str, output_path: Optional[str] = None, company_name: Optional[str] = None, job_title: Optional[str] = None) -> str:
    """
    Convert markdown CV to Word document with professional typography.
    Hybrid of Professional Classic + ATS Maximum: beautiful + searchable.

    Args:
        markdown_text: CV in markdown format
        output_path: Optional specific output path
        company_name: Company name for filename
        job_title: Job title for filename

    Returns:
        Path to generated .docx file
    """
    doc = Document()

    # Set page margins (1 inch sides, 0.75 inch top/bottom)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Helper functions for formatting
    def _add_bottom_border(paragraph):
        """Add a thin bottom border to a paragraph (ATS-safe)."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')  # 0.5pt
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')  # Black
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _set_paragraph_background(paragraph, hex_color):
        """Set paragraph background color (shading). hex_color is a 6-char hex string like '6B46C1'."""
        pPr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color.upper())
        pPr.append(shd)

    def _add_name(text):
        """Add candidate name as H1."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(18)
        return p

    def _add_contact_line():
        """Add contact info from env vars."""
        parts = []
        if MAYTE_EMAIL:
            parts.append(MAYTE_EMAIL)
        if MAYTE_PHONE:
            parts.append(MAYTE_PHONE)
        parts.append("Ciudad de Mexico, Mexico")
        if MAYTE_LINKEDIN:
            parts.append(MAYTE_LINKEDIN)
        if MAYTE_GITHUB:
            parts.append(MAYTE_GITHUB)

        if not parts:
            return

        contact_text = " | ".join(parts)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(8)
        run = p.add_run(contact_text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

    def _add_section_header(text):
        """Add section header with purple background bar and white text."""
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(10)
        pf.space_after = Pt(4)
        pf.left_indent = Inches(0.1)
        _set_paragraph_background(p, '6B46C1')
        run = p.add_run(' ' + text.upper())
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        return p

    def _add_subsection(text):
        """Add subsection header (company/project name) in purple."""
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(1)
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x6B, 0x46, 0xC1)
        return p

    def _add_subtitle(text):
        """Add subtitle (professional title) in purple."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(1)
        pf.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x6B, 0x46, 0xC1)
        return p

    def _add_body(text):
        """Add body paragraph with inline formatting."""
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(3)
        pf.line_spacing = Pt(14)  # ~1.15 line spacing

        # Handle inline bold (**text**)
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
            else:
                run = p.add_run(part)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)

    def _add_bullet(text):
        """Add bullet point."""
        p = doc.add_paragraph(text, style='List Bullet')
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        pf.left_indent = Inches(0.25)
        pf.first_line_indent = Inches(-0.15)
        # Format font
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

    # Parse and convert markdown
    lines = markdown_text.split('\n')
    first_heading = True
    current_list = []

    def flush_list():
        """Add accumulated list items to document."""
        nonlocal current_list
        if current_list:
            for item in current_list:
                _add_bullet(item)
            current_list = []

    for line in lines:
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            flush_list()
            continue

        # H1 (Name)
        if stripped.startswith('# '):
            flush_list()
            _add_name(stripped[2:])
            if first_heading:
                _add_contact_line()
                first_heading = False
            continue

        # Italic subtitle (professional title)
        if (stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**')) or \
           (stripped.startswith('_') and stripped.endswith('_')):
            flush_list()
            text = stripped.lstrip('*_').rstrip('*_')
            _add_subtitle(text)
            continue

        # H2 (Section header)
        if stripped.startswith('## '):
            flush_list()
            _add_section_header(stripped[3:])
            continue

        # H3 (Subsection)
        if stripped.startswith('### '):
            flush_list()
            _add_subsection(stripped[4:])
            continue

        # List items
        if stripped.startswith('- '):
            current_list.append(stripped[2:])
            continue

        # Regular text / paragraphs
        flush_list()
        _add_body(stripped)

    # Flush any remaining list items
    flush_list()

    # Generate output path
    if output_path is None:
        output_dir = Path("./cv_output")
        output_dir.mkdir(exist_ok=True)

        # Create safe filename with company and job title
        safe_company = re.sub(r'[^\w\s-]', '', company_name or "CV") if company_name else "CV"
        safe_company = safe_company.strip().replace(' ', '_')

        safe_title = re.sub(r'[^\w\s-]', '', job_title or "") if job_title else ""
        safe_title = safe_title.strip().replace(' ', '_')[:40]

        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        if safe_title:
            filename = f"{safe_company}_{safe_title}_CV_{timestamp}.docx"
        else:
            filename = f"{safe_company}_CV_{timestamp}.docx"

        output_path = output_dir / filename

    # Save document
    doc.save(str(output_path))
    return str(output_path)
